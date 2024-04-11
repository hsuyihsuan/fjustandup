from fastapi import Request, Depends, Form
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from docx import Document
import crud, function, schemas
from database import SessionLocal
from sqlalchemy.orm import Session
from datetime import date
from fastapi import APIRouter
from fastapi.responses import FileResponse
from docx2pdf import convert



router = APIRouter()

templates = Jinja2Templates(directory="templates") # 為能前後端分離, 將前端介面存在templates資料夾裡供使用, jinja2是知名樣板(template)處理器

def get_db(): #給予資料庫session啟用
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()







#查看病歷頁面
@router.post("/{patient_id}_{visit_date}_report", response_class=FileResponse, tags=["watch"])
async def download(
   request: Request, 
   patient_id: str,
   visit_date: str, 
   RecordType: str = Form(),
   Reason: str = Form(),
   db: Session = Depends(get_db)
): 
    session_now = request.session.get("user_role", None)
    if session_now in ['doctor']:
        today = date.today()

        crud.create_active_record(db, RecordType, patient_id, visit_date, Reason, request.session['account_id'])
        patient = crud.get_patient(db, patient_id)
        patient.PatientName = function.decrypt_data(patient.PatientName, patient.PatientKey)
        visit = crud.get_patient_visit_by_date(db, patient_id, visit_date)
        
        document = Document()

        document.add_heading('病歷資訊報表', 0)
        document.add_paragraph('列印日期: '+str(today))
        document.add_paragraph('就診日期: '+str(visit_date))
        document.add_paragraph('下載醫師: '+str(crud.get_account_user_name_by_account_id(db, request.session['account_id']).UserName))

        document.add_heading('個案資料卡', level=1)
        table = document.add_table(rows=8, cols=6, style="Table Grid")

        def make_table_heading_bold(row, heading):
             row[0].text = heading
             run = row[0].paragraphs[0].runs
             font = run[0].font
             font.bold = True
             return

        make_table_heading_bold(table.rows[0].cells, "簡歷")

        def merge_col(start, end):
            start.merge(end)
            return

        merge_col(table.cell(0, 0), table.cell(0, 5))

        row = table.rows[1].cells
        row[0].text = "病歷號碼"
        row[1].text = patient_id
        row[2].text = "姓名"
        row[3].text = patient.PatientName
        row[4].text = "性別"
        if patient.PatientSex:
            row[5].text = "男"
        else:
            row[5].text = "女"

        row = table.rows[2].cells
        row[0].text = "出生日期"
        row[1].text = str(patient.PatientBirth)
        row[2].text = "年紀"
        row[3].text = str(function.calculate_age(patient.PatientBirth, today))+"歲"
        row[4].text = "初始症狀"
        if patient.BeginSymptom == 0:
            row[5].text = "眼肌型"
        elif patient.BeginSymptom == 1:
            row[5].text = "口咽型"
        else:
            row[5].text = "四肢型"

        merge_col(table.cell(3, 0), table.cell(3, 5))

        make_table_heading_bold(table.rows[4].cells, "發病相關")
        merge_col(table.cell(4, 0), table.cell(4, 5))

        row = table.rows[5].cells
        row[0].text = "發病日期"
        row[1].text = str(patient.AttackDate)
        row[2].text = "發病年紀"
        row[3].text = str(function.calculate_age(patient.PatientBirth, patient.AttackDate))+"歲"
        row[4].text = "發病時長"
        row[5].text = str(function.days_of_onset(patient.AttackDate, today))+"天"

        row = table.rows[6].cells
        row[0].text = "最近本院住院日期"
        row[3].text = "本院住院總次數"

        try:
            row[2].text = str(crud.get_newest_in_hospital_date_by_patient_id(db, patient_id).VisitDate)
            row[5].text = str(crud.get_in_hospital_count(db, patient_id))
        except:
            row[2].text = "尚無紀錄"
            row[5].text = "0"

       

        merge_col(table.cell(6, 0), table.cell(6, 1))
        merge_col(table.cell(6, 3), table.cell(6, 4))

        row = table.rows[7].cells
        row[0].text = "最近外院住院日期"

        if str(patient.InOtherHospitalDate) == "0001-01-01":
            row[2].text = "尚無紀錄"
        else:
            row[2].text = str(patient.InOtherHospitalDate)
             
        row[3].text = "外院住院總次數"
        row[5].text = str(patient.InOtherHospitalCount)

        merge_col(table.cell(7, 0), table.cell(7, 1))
        merge_col(table.cell(7, 3), table.cell(7, 4))

        document.add_heading('本次就診資訊', level=1)
        table = document.add_table(rows=16, cols=6, style="Table Grid")

        make_table_heading_bold(table.rows[0].cells, "基本資訊")
        merge_col(table.cell(0, 0), table.cell(0, 5))

        row = table.rows[1].cells
        row[0].text = "就診狀態"
        if visit.Treat == 4:
            row[1].text = "開刀"
        elif visit.Treat == 3:
            row[1].text = "RTX"
        elif visit.Treat == 2:
            row[1].text = "IVIG"
        else:
            row[1].text = "看診"

        row[3].text = "MGFA臨床分類"
        if visit.MGFA_Classification == 0:
            row[5].text = "I"
        elif visit.MGFA_Classification == 1:
            row[5].text = "IIA"
        elif visit.MGFA_Classification== 2:
            row[5].text = "IIB"
        elif visit.MGFA_Classification == 3:
            row[5].text = "IIIA"
        elif visit.MGFA_Classification == 4:
            row[5].text = "IIIB"
        elif visit.MGFA_Classification == 5:
            row[5].text = "IVA"
        elif visit.MGFA_Classification == 6:
            row[5].text = "IVB"
        else:
            row[5].text = "V"

        merge_col(table.cell(1, 1), table.cell(1, 2))
        merge_col(table.cell(1, 3), table.cell(1, 4))

        row = table.rows[2].cells
        row[0].text = "身高"
        row[1].text = str(round(visit.Height, 1))
        row[2].text = "體重"
        row[3].text = str(round(visit.Weight, 1))
        row[4].text = "BMI"
        row[5].text = str(function.bmi(visit.Height, visit.Weight))

        row = table.rows[3].cells
        row[0].text = "收縮壓"
        if visit.Sbp > -1:
            row[1].text = str(visit.Sbp)
        else:
            row[1].text  = "尚無紀錄"
        
        row[2].text = "舒張壓"
        if visit.Dbp > -1:
            row[3].text = str(visit.Dbp)
        else:
            row[3].text = "尚無紀錄"

        merge_col(table.cell(3, 4), table.cell(3, 5))

        row = table.rows[4].cells
        row[0].text = "現有症狀"
        symptom_list = []

        if visit.Ptosis:
            symptom_list.append("眼瞼下垂")
        if visit.Diplopia:
            symptom_list.append("複視")
        if visit.Dysphasia:
            symptom_list.append("吞嚥困難")
        if visit.Dysarthria:
            symptom_list.append("講話不清")
        if visit.Dyspnea:
            symptom_list.append("呼吸困難")
        if visit.LimbWeakness:
            symptom_list.append("手腳無力")
        
        if symptom_list:
            table_cell = symptom_list[0]
            for i in range(1, len(symptom_list)):
                table_cell += "、"+symptom_list[i]

            row[1].text = table_cell
        else:
            row[1].text = "尚無症狀"

        merge_col(table.cell(4, 1), table.cell(4, 5))

        row = table.rows[5].cells
        row[0].text = "其他疾病"
        other_disease = crud.get_patient_other_disease_by_visit_date(db, patient_id, visit_date)

        # if other_disease:
        #     table_cell = other_disease[0].OtherDiseaseName
        #     for i in range(1, other_disease.count()):
        #         table_cell += "、"+other_disease[i].OtherDiseaseName
            
        #     row[1].text = table_cell
        # else:
        #     row[1].text = "尚無其他疾病"


        merge_col(table.cell(5, 1), table.cell(5, 5))

        row = table.rows[6].cells
        row[0].text = "診斷醫師"
        row[1].text = str(crud.get_visit_input_account(db, visit.PK_VisitID, visit.FK_AccountID).UserName)

        merge_col(table.cell(6, 1), table.cell(6, 2))
        merge_col(table.cell(6, 3), table.cell(6, 5))
        merge_col(table.cell(7, 0), table.cell(7, 5))

        make_table_heading_bold(table.rows[8].cells, "用藥資訊(每週頻率)")

        merge_col(table.cell(8, 0), table.cell(8, 5))

        row = table.rows[9].cells
        row[0].text = "大力丸(Pyridostigmine)"
        row[2].text = str(visit.Pyridostigmine) + "顆"

        row[3].text = "類固醇(Compesolone)"
        row[5].text = str(visit.Compesolone) + "顆"

        merge_col(table.cell(9, 0), table.cell(9, 1))
        merge_col(table.cell(9, 3), table.cell(9, 4))

        row = table.rows[10].cells
        row[0].text = "山喜多(Cellcept)"
        row[2].text = str(visit.Cellcept) + "顆"

        row[3].text = "移護寧(Imuran)"
        row[5].text = str(visit.Imuran) + "顆"

        merge_col(table.cell(10, 0), table.cell(10, 1))
        merge_col(table.cell(10, 3), table.cell(10, 4))

        row = table.rows[11].cells
        row[0].text = "普洛可富(Prograf)"
        row[2].text = str(visit.Prograf) + "顆"

        merge_col(table.cell(11, 0), table.cell(11, 1))
        merge_col(table.cell(11, 3), table.cell(11, 5))

        row = table.rows[12].cells
        row[0].text = "其他用藥"
        if other_medicine := crud.get_patient_other_medicine_by_visit_date(db, patient_id, visit_date):
            table_cell = other_medicine[0].OtherMedicineName + ": " + str(other_medicine[0].OtherMedicineCount)
            for i in range(1, len(other_medicine)):
                table_cell += "、" + other_medicine[i].OtherMedicineName + ": " + str(other_medicine[i].OtherMedicineCount) + "顆"
            row[1].text = table_cell
        else:
            row[1].text = "尚無其他用藥"

        

        merge_col(table.cell(12, 1), table.cell(12, 5))
        merge_col(table.cell(13, 0), table.cell(13, 5))

        make_table_heading_bold(table.rows[0].cells, "評估資訊")

        merge_col(table.cell(13, 0), table.cell(13, 5))

        row = table.rows[14].cells
        row[0].text = "自覺嚴重程度"
        if visit.SelfAssessment == 0:
            row[2].text = "輕度"
        elif visit.SelfAssessment == 1:
            row[2].text = "中度"
        else:
            row[2].text = "重度"

        merge_col(table.cell(14, 0), table.cell(14, 1))
        merge_col(table.cell(14, 2), table.cell(14, 5))


        row = table.rows[15].cells
        row[0].text = "附註"
        if visit.Note == "na":
            row[1].text = "尚無附註"
        else:
            row[1].text = str(visit.Note)


        merge_col(table.cell(15, 1), table.cell(15, 5))

        document.add_page_break()

        document.add_heading('胸腺斷層掃描紀錄', level=1)


        if thymus := crud.get_patient_all_thymus(db, patient_id):

            table = document.add_table(rows=len(thymus)+1, cols=6, style="Table Grid")
            row = table.rows[0].cells
            row[0].text = "斷層日期"
            row[1].text = "斷層結果"
            row[2].text = "病理描述"
            merge_col(table.cell(0, 2), table.cell(0, 5))

            for i in range(len(thymus)):
                row = table.rows[i+1].cells
                row[0].text = str(thymus[i].CtDate)
                if thymus[i].ThymusStatus == 0:
                    row[1].text = "正常"
                elif thymus[i].ThymusStatus == 1:
                    row[1].text = "胸腺萎縮"
                elif thymus[i].ThymusStatus == 2:
                    row[1].text = "胸腺增生"
                else:
                    row[1].text = "胸腺瘤"
                
                if thymus[i].ThymusDescription == "na":
                    row[2].text = "尚無附註"
                else:
                    row[2].text = str(thymus[i].ThymusDescription)

                merge_col(table.cell(i+1, 2), table.cell(i+1, 5))

           
        else:
            table = document.add_table(rows=2, cols=6, style="Table Grid")
            row = table.rows[0].cells
            row[0].text = "斷層日期"
            row[1].text = "斷層結果"
            row[2].text = "病理描述"
            merge_col(table.cell(0, 2), table.cell(0, 5))
            merge_col(table.cell(1, 2), table.cell(1, 5))


        document.add_heading('抗體抽血紀錄', level=1)


        if blood_test := crud.get_patient_all_blood_test(db, patient_id):
            table = document.add_table(rows=len(blood_test)+1, cols=6, style="Table Grid")

            row = table.rows[0].cells
            row[0].text = "抽血日期"
            row[1].text = "AchR抗體"
            row[2].text = "TSH"
            row[3].text = "fT4"
            row[4].text = "ANA"
            row[5].text = "Uric acid"


            for i in range(len(blood_test)):
                row = table.rows[i+1].cells
                row[0].text = str(blood_test[i].BloodTestDate)

                if blood_test[i].AchR == -1.0:
                    row[1].text = "無紀錄"
                else:
                    row[1].text = str(round(blood_test[i].AchR, 2))
                
                if blood_test[i].TSH == -1.0:
                    row[2].text = "無紀錄"
                else:
                    row[2].text = str(round(blood_test[i].TSH, 2))

                if blood_test[i].FreeThyroxine == -1.0:
                    row[3].text = "無紀錄"
                else:
                    row[3].text = str(round(blood_test[i].FreeThyroxine, 2))

                if blood_test[i].ANA == -1.0:
                    row[4].text = "無紀錄"
                else:
                    row[4].text = str(blood_test[i].ANA)

                if blood_test[i].UricAcid == -1.0:
                    row[5].text = "無紀錄"
                else:
                    row[5].text = str(round(blood_test[i].UricAcid, 2))

        else:
            table = document.add_table(rows=2, cols=6, style="Table Grid")

            row = table.rows[0].cells
            row[0].text = "抽血日期"
            row[1].text = "AchR抗體"
            row[2].text = "TSH"
            row[3].text = "fT4"
            row[4].text = "ANA"
            row[5].text = "Uric acid"


                
                

        # document.add_heading('量表紀錄', level=1)
        # table = document.add_table(rows=8, cols=8, style="Table Grid")

        # row = table.rows[0].cells
        # row[0].text = "QOL"
        # row[1].text = "總分"
        # row[2].text = "QMG Score"
        # row[3].text = "總分"
        # row[4].text = "MG Composite"
        # row[5].text = "總分"
        # row[6].text = "ADL"
        # row[7].text = "總分"

        document.save('report.docx')

        file_path = "report.docx"

        convert("report.docx")

        headers = {'Content-Disposition': 'attachment; filename="report.pdf"'}

        path = "report.pdf"

  

        return FileResponse(path, headers=headers)

















        
        # today = date.today()
        # patient = crud.get_patient(db, patient_id)
        # patient_age = function.calculate_age(patient.PatientBirth, today)
        # visit = crud.get_patient_visit_by_date(db, patient_id, visit_date)

        # all_visit_date = []
        # for i in reversed(crud.get_all_visit_date_by_patient_id(db, patient_id)):
        #          all_visit_date.append(str(i.VisitDate))

        # days_of_onset = function.days_of_onset(patient.AttackDate, today)
        # age_of_onset = function.calculate_age(patient.PatientBirth, patient.AttackDate)
        # bmi = function.bmi(visit.Height, visit.Weight)

        # if in_hospital_date := crud.get_newest_in_hospital_date_by_patient_id(db, patient_id):
        #     in_hospital_date = in_hospital_date.VisitDate
        
        # in_hospital_count = crud.get_in_hospital_count(db, patient_id)

        # other_disease = crud.get_patient_other_disease_by_visit_date(db, patient_id, visit_date)
        # other_medicine = crud.get_patient_other_medicine_by_visit_date(db, patient_id, visit_date)

        # if first_time_thymus_atrophy := crud.get_patient_first_time_thymus_atrophy(db, patient_id):
        #     first_time_thymus_atrophy = first_time_thymus_atrophy.CtDate
        # if first_time_thymus_hyperplasia := crud.get_patient_first_time_thymus_hyperplasia(db, patient_id):
        #     first_time_thymus_hyperplasia = first_time_thymus_hyperplasia.CtDate
        # if first_time_thymus_thymoma := crud.get_patient_first_time_thymus_thymoma(db, patient_id):
        #     first_time_thymus_thymoma = first_time_thymus_thymoma.CtDate

        # all_thymus_dates = []
        # one_thymus = None
        # thymus_date = None
        # if thymus := crud.get_patient_all_thymus(db, patient_id):
        #     for i in thymus:
        #         all_thymus_dates.append(str(i.CtDate))
        #     one_thymus = crud.get_patient_newest_thymus(db, patient_id)
        #     thymus_date = one_thymus.CtDate
            

            

        # all_blood_test_dates = []
        # one_blood_test = None
        # blood_test_date = None
        # AchR = []
        # TSH = []
        # ANA = []
        # if blood_test := crud.get_patient_all_blood_test(db, patient_id):
        #     for i in blood_test:
        #         all_blood_test_dates.append(str(i.BloodTestDate))
        #         if i.AchR != -1:
        #             AchR.append(round(i.AchR, 2))
        #         else:
        #             AchR.append('None')
        #         if i.TSH != -1:
        #             TSH.append(round(i.TSH, 1))
        #         else:
        #             TSH.append('None')
        #         if i.ANA != -1:
        #              ANA.append(i.ANA)
        #         else:
        #              ANA.append('None')
                
          
               
            
        #     one_blood_test = crud.get_patient_newest_blood_test(db, patient_id)
        #     blood_test_date = one_blood_test.BloodTestDate

        # qol = None
        # qol_date = None
        # if qol_field_data := crud.get_patient_newest_qol(db, patient_id):
        #     qol_date = qol_field_data.QOL_Date
        #     qol = function.CreateTableFields()
        #     qol = qol.create_table_fields(qol.qol, qol_field_data)

        # #從這開始改
        # qmg = None
        # qmg_date = None
        # if qmg_field_data := crud.get_patient_newest_qmg(db, patient_id):
        #     qmg_date = qmg_field_data.QMG_Date
        #     qmg = function.CreateTableFields()
        #     qmg = qmg.create_table_fields(qmg.qmg, qmg_field_data)

        # mg_composite = None
        # mg_composite_date = None
        # if mg_composite_field_data := crud.get_patient_newest_mg_composite(db, patient_id):
        #     mg_composite_date = mg_composite_field_data.MGComposite_Date
        #     mg_composite = function.CreateTableFields()
        #     mg_composite = mg_composite.create_table_fields(mg_composite.mg_composite, mg_composite_field_data)

        # adl = None
        # adl_date = None
        # if adl_field_data := crud.get_patient_newest_adl(db, patient_id):
        #     adl_date = adl_field_data.ADL_Date
        #     adl = function.CreateTableFields()
        #     adl = adl.create_table_fields(adl.adl, adl_field_data)

        #     # if mg_composite_field_data:
        #     #     mg_composite_date = mg_composite_field_data.MGComposite_Date
        #     # else:
        #     #     mg_composite_date = None
        # # crud.get_patient_all_qol_date(db, patient_id)
        # all_qmg_dates = crud.get_patient_all_qmg_date(db, patient_id)
        # all_mg_composite_dates = crud.get_patient_all_mg_composite_date(db, patient_id)
        # all_adl_dates = crud.get_patient_all_adl_date(db, patient_id)
               
        # all_visit = crud.get_patient_all_visit(db, patient_id)
        
        # height = []
        # weight = []
        # sbp = []
        # dbp = []
        # pyridostigmine = []
        # compesolone = []
        # cellcept = []
        # imuran = []
        # prograf = []
        

        # for i in all_visit:
        #     height.append(round(i.Height, 1))
        #     weight.append(round(i.Weight, 1))
        #     sbp.append(i.Sbp)
        #     dbp.append(i.Dbp)
        #     pyridostigmine.append(i.Pyridostigmine)
        #     compesolone.append(i.Compesolone)
        #     cellcept.append(i.Cellcept)
        #     imuran.append(i.Imuran)
        #     prograf.append(i.Prograf)

        # all_qol = crud.get_patient_all_qol(db, patient_id)
        # all_qol_dates = []
        # qol_sum = []
        # for i in all_qol:
        #      all_qol_dates.append(str(i.QOL_Date))
        #      qol_sum.append(i.QOL_Sum)

        

        # all_qmg = crud.get_patient_all_qmg(db, patient_id)
        # all_qmg_dates = []
        # qmg_sum = []
        # for i in all_qmg:
        #     all_qmg_dates.append(str(i.QMG_Date))
        #     qmg_sum.append(i.QMG_Sum)

        # all_mg_composite = crud.get_patient_all_mg_composite(db, patient_id)
        # all_mg_composite_dates = []
        # mg_composite_sum = []
        # for i in all_mg_composite:
        #     all_mg_composite_dates.append(str(i.MGComposite_Date))
        #     mg_composite_sum.append(i.MGComposite_Sum)

        # all_adl = crud.get_patient_all_adl(db, patient_id)
        # all_adl_dates = []
        # adl_sum = []
        # for i in all_adl:
        #     all_adl_dates.append(str(i.ADL_Date))
        #     adl_sum.append(i.ADL_Sum)

        
        # try:
        #     model_url = crud.get_model_which_is_activate(db).PK_CreateTime
        # except:
        #     model_url = None


        # return templates.TemplateResponse("watchVisit.html", {"request": request, 
        #                                                       "patient": patient, 
        #                                                       "patient_id": patient_id, 
        #                                                       "patient_age": patient_age, 
        #                                                       "visit": visit, 
        #                                                       "all_visit_date": all_visit_date, 
        #                                                       "all_visit_date_asc": list(reversed(all_visit_date)), 
        #                                                       "days_of_onset": days_of_onset, 
        #                                                       "age_of_onset": age_of_onset,
        #                                                       "bmi": bmi,
        #                                                       "in_hospital_date": in_hospital_date,
        #                                                       "in_hospital_count": in_hospital_count,
        #                                                       "other_disease": other_disease,
        #                                                       "other_medicine": other_medicine,
        #                                                       "thymus": thymus, 
        #                                                       "all_thymus_dates": all_thymus_dates, 
        #                                                       "one_thymus": one_thymus,
        #                                                       "thymus_date": thymus_date,
        #                                                       "first_time_thymus_atrophy": first_time_thymus_atrophy, 
        #                                                       "first_time_thymus_hyperplasia": first_time_thymus_hyperplasia, 
        #                                                       "first_time_thymus_thymoma": first_time_thymus_thymoma,
        #                                                       "blood_test": blood_test, 
        #                                                       "all_blood_test_dates": all_blood_test_dates, 
        #                                                     #   "all_blood_test_dates_asc": list(reversed(all_blood_test_dates)), 
        #                                                       "one_blood_test": one_blood_test,
        #                                                     #   "blood_test_date": blood_test_date,
        #                                                       "AchR": AchR,
        #                                                       "TSH": TSH,
        #                                                       "ANA": ANA,

        #                                                       "all_qol_dates": all_qol_dates,
        #                                                       "qol": qol,
        #                                                       "qol_sum": qol_sum,
        #                                                       "qol_date": qol_date,

        #                                                       "all_qmg_dates": all_qmg_dates,
        #                                                       "qmg": qmg,
        #                                                       "qmg_sum": qmg_sum,
        #                                                       "qmg_date": qmg_date,

        #                                                       "all_mg_composite_dates": all_mg_composite_dates,
        #                                                       "mg_composite": mg_composite, 
        #                                                       "mg_composite_sum": mg_composite_sum,
        #                                                       "mg_composite_date": mg_composite_date,
                                                              
        #                                                       "all_adl_dates": all_adl_dates,
        #                                                       "adl": adl,
        #                                                       "adl_sum": adl_sum,
        #                                                       "adl_date": adl_date,

        #                                                       "watch": 1,
        #                                                       "basic": "height_and_weight",
        #                                                     #   "height": height,
        #                                                     #   "weight": weight, 
        #                                                     #   "sbp": sbp,
        #                                                     #   "dbp": dbp,
        #                                                       "pyridostigmine": pyridostigmine,
        #                                                       "compesolone": compesolone,
        #                                                       "cellcept": cellcept, 
        #                                                       "imuran": imuran, 
        #                                                       "prograf": prograf, 
                                                          
        #                                                       "model_url": model_url,

        #                                                       })
    else:
         return HTMLResponse(status_code=403)